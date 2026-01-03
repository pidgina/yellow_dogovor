import csv
import random
import os
import re
from datetime import datetime, timedelta
from pathlib import Path

# Попытка импорта docx для создания профессиональных документов   123
try:
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

def sanitize_filename(name):
    """Удаляет символы, запрещенные в именах файлов."""
    if not name:
        return "Unknown"
    return re.sub(r'[\\/*?:"<>|]', "", str(name)).replace(" ", "_")

def get_random_doc_date(date_start_str=None):
    """
    Генерирует дату подписания договора.
    Если передана дата начала работ, договор подписывается в этот же день.
    """
    if date_start_str:
        return date_start_str
        
    today = datetime.now()
    # Примерно 150 дней в 5 месяцах
    random_days_ago = random.randint(0, 150)
    doc_date = today - timedelta(days=random_days_ago)
    return doc_date.strftime("%d.%m.%Y")

def create_docx_contract(data, filepath):
    """Создает профессиональный договор в формате .docx с таблицами"""
    doc = Document()

    # Настройка стиля по умолчанию (Times New Roman, если есть)
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)

    # Заголовок
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(f"ДОГОВОР № {data['idx']}")
    run.bold = True
    run.font.size = Pt(14)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run("об оказании возмездных услуг").italic = True

    # Город и дата
    city_date = doc.add_paragraph()
    city_date.add_run(f"г. Москва")
    # Табуляция для выравнивания даты вправо
    city_date.add_run(f"\t\t\t\t\t\t\t\t{data['doc_date']}")

    # Преамбула
    p = doc.add_paragraph()
    p.add_run(f"Гражданин(ка) ")
    p.add_run(data['zakazchik']).bold = True
    p.add_run(", именуемый(ая) в дальнейшем \"Заказчик\", с одной стороны, и Гражданин(ка) ")
    p.add_run(data['ispolnitel']).bold = True
    p.add_run(", именуемый(ая) в дальнейшем \"Исполнитель\", с другой стороны, вместе именуемые \"Стороны\", заключили настоящий договор о нижеследующем:")

    # Разделы
    doc.add_heading("1. ПРЕДМЕТ ДОГОВОРА И СТОИМОСТЬ", level=2)
    doc.add_paragraph("1.1. Исполнитель обязуется оказать услуги, а Заказчик обязуется принять и оплатить их в соответствии с условиями настоящего договора:")
    
    # ТАБЛИЦА С ДЕТАЛЯМИ УСЛУГ
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Наименование услуги'
    hdr_cells[1].text = 'С даты'
    hdr_cells[2].text = 'По дату'
    hdr_cells[3].text = 'Итого'
    
    # Делаем заголовки жирными
    for cell in hdr_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True

    row_cells = table.add_row().cells
    row_cells[0].text = data['uslugi']
    row_cells[1].text = data['date_start']
    row_cells[2].text = data['date_end']
    row_cells[3].text = f"{data['stoimost']} руб."

    doc.add_heading("2. ПОРЯДОК ОПЛАТЫ", level=2)
    doc.add_paragraph("2.1. Оплата услуг производится Заказчиком в размере 100% предоплаты до начала оказания услуг.")

    # ПОДПИСИ СТОРОН
    doc.add_paragraph("\n")
    doc.add_heading("РЕКВИЗИТЫ И ПОДПИСИ СТОРОН", level=2)
    
    sig_table = doc.add_table(rows=2, cols=2)
    sig_cells_top = sig_table.rows[0].cells
    sig_cells_top[0].text = "ЗАКАЗЧИК"
    sig_cells_top[1].text = "ИСПОЛНИТЕЛЬ"
    
    for cell in sig_cells_top:
        for paragraph in cell.paragraphs:
            if paragraph.runs:
                paragraph.runs[0].bold = True

    sig_cells_bot = sig_table.rows[1].cells
    sig_cells_bot[0].text = f"\n________________ / {data['zakazchik']} /\n\n"
    sig_cells_bot[1].text = f"\n________________ / {data['ispolnitel']} /\n\n"

    doc.save(filepath)

def create_txt_fallback(data, filepath):
    """Создает текстовый договор (если docx недоступен)"""
    template = """
================================================================================
                                ДОГОВОР № {idx}
                об оказании услуг между физическими лицами

г. Москва                                                     {doc_date}

Гражданин(ка) {zakazchik}, именуемый(ая) в дальнейшем "Заказчик",
и
Гражданин(ка) {ispolnitel}, именуемый(ая) в дальнейшем "Исполнитель",

заключили настоящий договор о нижеследующем:

1. ПРЕДМЕТ ДОГОВОРА
1.1. Исполнитель обязуется оказать следующие услуги: {uslugi}.
1.2. Срок оказания услуг: с "{date_start}" по "{date_end}".

2. СТОИМОСТЬ И ПОРЯДОК ОПЛАТЫ
2.1. Стоимость услуг составляет {stoimost} рублей за весь объём работ.
2.2. Предоплата до начала работ в размере 100% от итоговой суммы.

РЕКВИЗИТЫ И ПОДПИСИ СТОРОН:
Заказчик: ________________ /{zakazchik}/
Исполнитель: ________________ /{ispolnitel}/
================================================================================
"""
    with open(filepath, "w", encoding="utf-8") as f:
        f.write(template.format(**data))

def get_next_contract_number(output_path):
    """Находит максимальный номер договора в папке и возвращает следующий."""
    max_num = 0
    if not output_path.exists():
        return 1
    
    # Регулярное выражение для поиска номера в начале файла: Договор_123_...
    pattern = re.compile(r"Договор_(\d+)_")
    
    for file in output_path.iterdir():
        match = pattern.match(file.name)
        if match:
            num = int(match.group(1))
            if num > max_num:
                max_num = num
    return max_num + 1

def format_fio(name):
    """
    Форматирует ФИО: каждая часть с большой буквы. 
    Разрешает дефисы внутри слов.
    Возвращает (отформатированное имя, количество слов).
    """
    if not name:
        return "", 0
    
    # Очистка от лишних пробелов и разделение на слова
    # Регулярное выражение находит слова, разделенные пробелами или дефисами
    parts = re.split(r'(\s+|-)', name.strip())
    
    formatted_parts = []
    word_count = 0
    
    for part in parts:
        if part.strip() and part != '-':
            # Это слово - делаем заглавную букву
            formatted_parts.append(part.capitalize())
            word_count += 1
        else:
            # Это разделитель (пробел или дефис)
            formatted_parts.append(part)
            
    return "".join(formatted_parts).strip(), word_count

def batch_contracts(csv_file, output_dir="ready_contracts"):
    output_path = Path(output_dir)
    output_path.mkdir(exist_ok=True)

    print(f"--- Генератор профессиональных договоров ---")
    
    # Пытаемся импортировать генератор для авто-обновления
    try:
        from generator_clientov import generate_fake_data
        can_refresh = True
    except ImportError:
        can_refresh = False

    # Первичная проверка/генерация файла
    if not os.path.exists(csv_file):
        if can_refresh:
            print("Файл данных не найден. Генерирую начальный список...")
            generate_fake_data(csv_file, 50)
        else:
            print(f"Ошибка: Файл {csv_file} не найден и генератор недоступен!")
            return

    global_ispolnitel = ""
    while True:
        raw_ispolnitel = input("Введите ФИО исполнителя (пусто - из файла): ").strip()
        if not raw_ispolnitel:
            break
            
        global_ispolnitel, word_count = format_fio(raw_ispolnitel)
        
        if word_count > 3:
            print(f"Ошибка: Введено {word_count} слов, а разрешено максимум 3. Попробуйте еще раз.")
        else:
            if global_ispolnitel != raw_ispolnitel:
                print(f"Имя отформатировано: {global_ispolnitel}")
            break

    # Спрашиваем лимит
    try:
        limit_input = input("Сколько новых договоров создать? (например, 150): ").strip()
        limit = int(limit_input) if limit_input else 50
        if limit <= 0:
            print("Отмена.")
            return
    except ValueError:
        limit = 50

    # Определяем стартовый номер для нумерации
    next_idx = get_next_contract_number(output_path)
    count = 0
    current_rows = []
    row_in_batch_idx = 0
    
    while count < limit:
        # ОБНОВЛЕНИЕ СПИСКА: Каждые 50 договоров (или в начале)
        if count % 50 == 0:
            if can_refresh:
                if count > 0:
                    print("\nОбновляю список заказчиков...")
                generate_fake_data(csv_file, 50)
            
            with open(csv_file, encoding='utf-8') as f:
                current_rows = list(csv.DictReader(f))
            row_in_batch_idx = 0

        # Берем строку из текущего набора (50 шт)
        if row_in_batch_idx >= len(current_rows):
            # На случай если в файле меньше 50 строк
            row_in_batch_idx = 0

        row = current_rows[row_in_batch_idx]
        
        ispolnitel = global_ispolnitel if global_ispolnitel else row.get('ispolnitel', 'Не указан')
        
        # Дата договора = За 0-3 дня до даты начала работ
        date_start = row.get('date_start', '__________')
        doc_date = get_random_doc_date(date_start)
        
        data = {
            'idx': next_idx,
            'zakazchik': row.get('zakazchik', '__________'),
            'ispolnitel': ispolnitel,
            'uslugi': row.get('uslugi', '__________'),
            'date_start': date_start,
            'date_end': row.get('date_end', '__________'),
            'stoimost': row.get('stoimost', '0'),
            'doc_date': doc_date
        }
        
        clean_zakazchik = sanitize_filename(data['zakazchik'])
        
        if HAS_DOCX:
            fname = f"Договор_{next_idx}_{clean_zakazchik}.docx"
            create_docx_contract(data, output_path / fname)
        else:
            fname = f"Договор_{next_idx}_{clean_zakazchik}.txt"
            create_txt_fallback(data, output_path / fname)
        
        next_idx += 1
        count += 1
        row_in_batch_idx += 1
        
        if count % 10 == 0 or count == limit:
            print(f"Готово: {count}/{limit} (Договор №{next_idx-1})...")

    fmt = "DOCX" if HAS_DOCX else "TXT"
    print(f"\nУспех! Сгенерировано {count} договоров в формате {fmt}.")
    print(f"Нумерация завершена на номере: {next_idx - 1}")
    print(f"Папка: {output_path.absolute()}")

if __name__ == "__main__":
    batch_contracts("contracts.csv")
